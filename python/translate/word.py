import threading
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
import translate
import common
import os
import sys
import time
import datetime

def start(trans):
    # 允许的最大线程
    threads=trans['threads']
    if threads is None or threads=="" or int(threads)<0:
        max_threads=10
    else:
        max_threads=int(threads)
    # 当前执行的索引位置
    run_index=0
    max_chars=1000
    start_time = datetime.datetime.now()
    # 创建Document对象，加载Word文件
    try:
        document = Document(trans['file_path'])
    except Exception as e:
        translate.error(trans['id'],trans['process_file'], "无法访问该文档")
        return False
    texts=[]

    # 仅文字-保留原文-重排
    # if trans['type']=='trans_text_both_new':

    # read_paragraph_text(document, texts)
    read_rune_text(document, texts)

    # print(texts)
    # exit()
    max_run=max_threads if len(texts)>max_threads else len(texts)
    event=threading.Event()
    before_active_count=threading.activeCount()
    while run_index<=len(texts)-1:
        if threading.activeCount()<max_run+before_active_count:
            if not event.is_set():
                thread = threading.Thread(target=translate.get,args=(trans,event,texts,run_index))
                thread.start()
                run_index+=1
            else:
                return False
    
    while True:
        if event.is_set():
            return False
        complete=True
        for text in texts:
            if not text['complete']:
                complete=False
        if complete:
            break
        else:
            time.sleep(1)
    print(texts)
    print("翻译文本-结束")
    # text_count=write_rune_text(document, texts)
    text_count=write_rune_both(document, texts)
    # text_count=write_paragraph_text(document, texts)


    # print("编辑文档-结束")
    # print(datetime.datetime.now())
    document.save(trans['target_file'])
    end_time = datetime.datetime.now()
    spend_time=common.display_spend(start_time, end_time)
    translate.complete(trans,text_count,spend_time)
    return True


def read_paragraph_text(document, texts):
    for paragraph in document.paragraphs:
        append_text(paragraph.text, texts)

def write_paragraph_text(document, texts):
    for paragraph in document.paragraphs:
        text=paragraph.text
        if check_text(text) and len(texts)>0:
            item=texts.pop(0)
            paragraph.text=item.get('text',"")
            # paragraph.text=paragraph.text+"\n"+item.get('text',"")

def read_rune_text(document, texts):
    for paragraph in document.paragraphs:
        read_run(paragraph.runs, texts)
        
        if len(paragraph.hyperlinks)>0:
            for hyperlink in paragraph.hyperlinks:
                read_run(hyperlink.runs, texts)

    # print("翻译文本--开始")
    # print(datetime.datetime.now())
    for table in document.tables:
        for row in table.rows:
            start_span=0
            for cell in row.cells:
                start_span+=1
                # if start_span==cell.grid_span:
                #     start_span=0
                    # read_cell(cell, texts)
                for index,paragraph in enumerate(cell.paragraphs):
                    # print(index)
                    # print(paragraph.text)
                    read_run(paragraph.runs, texts)

                    if len(paragraph.hyperlinks)>0:
                        for hyperlink in paragraph.hyperlinks:
                            read_run(hyperlink.runs, texts)

def write_rune_text(document, texts):
    text_count=0
    for paragraph in document.paragraphs:
        text_count+=write_run(paragraph.runs, texts)

        if len(paragraph.hyperlinks)>0:
            for hyperlink in paragraph.hyperlinks:
                text_count+=write_run(hyperlink.runs, texts)

    for table in document.tables:
        for row in table.rows:
            start_span=0
            for cell in row.cells:
                # start_span+=1
                # if start_span==cell.grid_span:
                #     start_span=0
                    # text_count+=write_cell(cell, texts)
                for paragraph in cell.paragraphs:
                    text_count+=write_run(paragraph.runs, texts)

                    if len(paragraph.hyperlinks)>0:
                        for hyperlink in paragraph.hyperlinks:
                            text_count+=write_run(hyperlink.runs, texts)

#保留原译文
def write_rune_both(document, texts):
    text_count=0
    for paragraph in document.paragraphs:
        space_before=paragraph.paragraph_format.space_before
        space_after=paragraph.paragraph_format.space_after
        line_spacing=paragraph.paragraph_format.line_spacing
        print(paragraph.text)
        if(len(paragraph.runs)>0):
            paragraph.runs[-1].add_break()
            add_paragraph_run(paragraph, paragraph.runs, texts, text_count)
        if len(paragraph.hyperlinks)>0:
            paragraph.hyperlinks[0].runs[-1].add_break()
            for hyperlink in paragraph.hyperlinks:
                add_paragraph_run(paragraph, hyperlink.runs, texts, text_count)
        # if space_before!=None:
        #     paragraph.paragraph_format.space_before=Inches(space_before).pt
        # if space_after!=None:
        #     paragraph.paragraph_format.space_after=Inches(space_after).pt
        # if line_spacing!=None:
        #     paragraph.paragraph_format.line_spacing=line_spacing
             #   paragraph.add_run(hyperlink.text, hyperlink.runs[0].style)
        # text_count+=write_run(paragraph.runs, texts)
    for table in document.tables:
        for row in table.rows:
            # start_span=0
            for cell in row.cells:
                # start_span+=1
                # if start_span==cell.grid_span:
                #     start_span=0
                    # text_count+=write_cell(cell, texts)
                for paragraph in cell.paragraphs:
                    paragraph.runs[-1].add_break()
                    add_paragraph_run(paragraph, paragraph.runs, texts, text_count)

                    if len(paragraph.hyperlinks)>0:
                        for hyperlink in paragraph.hyperlinks:
                            add_paragraph_run(paragraph, hyperlink.runs, texts, text_count)


def read_run(runs,texts):
    # text=""
    if len(runs)>0 or len(texts)==0:
        for index,run in enumerate(runs):
            append_text(run.text, texts)
        #     if run.text=="":
        #         if len(text)>0 and not common.is_all_punc(text):        
        #             texts.append({"text":text, "complete":False})
        #             text=""
        #     else:
        #         text+=run.text
        # if len(text)>0 and not common.is_all_punc(text):
        #     texts.append({"text":text, "complete":False})

def append_text(text, texts):
    if check_text(text):        
        texts.append({"text":text, "complete":False})

def check_text(text):
    return len(text)>0 and not common.is_all_punc(text) 

def write_run(runs,texts):
    text_count=0
    if len(runs)==0:
        return text_count
    text=""
    for index,run in enumerate(runs):
        text=run.text
        if len(text)>0 and not common.is_all_punc(text) and len(texts)>0:
            item=texts.pop(0)
            text_count+=item.get('count',0)
            run.text=item.get('text',"")
        # if run.text=="":
        #     if len(text)>0 and not common.is_all_punc(text) and len(texts)>0:
        #         item=texts.pop(0)
        #         text_count+=item.get('count',0)
        #         runs[index-1].text=item.get('text',"")
        #         text=""
        # else:
        #     text+=run.text
        #     run.text=""
    # if len(text)>0 and not common.is_all_punc(text) and len(texts)>0:
    #     item=texts.pop(0)
    #     text_count+=item.get('count',0)
    #     runs[0].text=item.get('text',"")
    return text_count


def read_cell(cell,texts):
    append_text(cell.text, texts)


def write_cell(cell,texts):
    text=cell.text
    text_count=0
    if len(text)>0 and not common.is_all_punc(text) and len(texts)>0:
        item=texts.pop(0)
        text_count+=item.get('count',0)
        cell.text=item.get('text',"")
    return text_count

def add_paragraph_run(paragraph, runs, texts, text_count):
    for index,run in enumerate(runs):
        if check_text(run.text) and len(texts)>0:
            item=texts.pop(0)
            text_count+=item.get('count',0)
            new_run=paragraph.add_run(run.text, run.style)
            new_run.text=item.get('text',"")
            set_run_style(new_run, run)

def set_run_style(new_run, copy_run):
    new_run.font.italic= copy_run.font.italic
    new_run.font.strike= copy_run.font.strike
    new_run.font.bold= copy_run.font.bold
    new_run.font.size= copy_run.font.size
    new_run.font.color.rgb= copy_run.font.color.rgb
    new_run.underline= copy_run.underline
    new_run.style= copy_run.style
